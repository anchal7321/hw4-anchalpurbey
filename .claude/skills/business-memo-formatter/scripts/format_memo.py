"""
format_memo.py
Business Memo Formatter — core script for the business-memo-formatter skill.

Usage:
    python .claude/skills/business-memo-formatter/scripts/format_memo.py <input_file> --output-dir outputs

Outputs:
    outputs/formatted_memo.docx
    outputs/formatting_report.md
"""

import argparse
import os
import re
import sys
from datetime import datetime

# -- Dependency check ----------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("\nERROR: python-docx is not installed.")
    print("Please run:  py -m pip install python-docx\n")
    sys.exit(1)

# -- Constants -----------------------------------------------------------------
KNOWN_FIELDS = [
    "title", "to", "from", "date", "subject",
    "recommendation", "analysis", "risks", "conclusion"
]

# Header fields capture only the inline value on the same line as the label.
# Any following unlabeled lines are sent to additional_notes, not appended here.
HEADER_FIELDS = {"title", "to", "from", "date", "subject"}

# Aliases map alternative labels to a canonical KNOWN_FIELDS key.
# Keys and values are lowercase. Checked before KNOWN_FIELDS during parsing.
FIELD_ALIASES = {
    "re": "subject",
    "regarding": "subject",
}

REQUIRED_FIELDS = ["title", "to", "from", "date", "subject", "recommendation", "analysis", "conclusion"]
OPTIONAL_FIELDS = ["risks"]

WORD_COUNT_LIMIT = 750


# -- PARSING -------------------------------------------------------------------
def parse_memo(text: str) -> dict:
    """
    Parse labeled memo fields from the input text.
    Fields are detected case-insensitively.

    Alias resolution: 'Re:' and 'Regarding:' are treated as 'Subject:'.

    Any text that appears without a recognized label — whether before the first
    field, between fields, or after the last field — is collected into
    'additional_notes' so that no content is silently dropped.
    """
    sections = {field: "" for field in KNOWN_FIELDS}
    sections["additional_notes"] = ""

    current_field = None
    buffer = []
    unlabeled_buffer = []

    # Build a combined lookup: aliases + known fields, all lowercase
    # Each entry maps a label string to its canonical field key.
    all_labels = {}
    for alias, canonical in FIELD_ALIASES.items():
        all_labels[alias] = canonical
    for field in KNOWN_FIELDS:
        all_labels[field] = field

    lines = text.splitlines()

    for line in lines:
        matched = False
        for label, field in all_labels.items():
            pattern = re.compile(rf"^\s*{re.escape(label)}\s*:\s*(.*)", re.IGNORECASE)
            match = pattern.match(line)
            if match:
                # Flush previous buffer before switching fields
                if current_field:
                    existing = sections[current_field]
                    addition = "\n".join(buffer).strip()
                    sections[current_field] = (existing + "\n" + addition).strip() if existing else addition
                else:
                    # Unlabeled text before the first recognized field
                    unlabeled_buffer.extend(buffer)
                # Start new field
                current_field = field
                inline_value = match.group(1).strip()
                if field in HEADER_FIELDS:
                    # Header fields: capture inline value only; reset buffer to empty
                    # so any following unlabeled lines go to unlabeled_buffer, not here.
                    sections[current_field] = inline_value
                    current_field = None
                    buffer = []
                else:
                    buffer = [inline_value]
                matched = True
                break

        if not matched:
            if current_field:
                buffer.append(line)
            else:
                unlabeled_buffer.append(line)

    # Flush final buffer
    if current_field:
        existing = sections[current_field]
        addition = "\n".join(buffer).strip()
        sections[current_field] = (existing + "\n" + addition).strip() if existing else addition

    # Collect all unlabeled text into additional_notes
    if unlabeled_buffer:
        sections["additional_notes"] = "\n".join(unlabeled_buffer).strip()

    return sections


# -- VALIDATION ----------------------------------------------------------------
def validate_sections(sections: dict) -> dict:
    """
    Check which required fields are present or missing.
    Returns a dict with 'found', 'missing', and 'optional_missing' lists.
    """
    found = []
    missing = []
    optional_missing = []

    for field in REQUIRED_FIELDS:
        if sections.get(field, "").strip():
            found.append(field)
        else:
            missing.append(field)

    for field in OPTIONAL_FIELDS:
        if not sections.get(field, "").strip():
            optional_missing.append(field)

    return {"found": found, "missing": missing, "optional_missing": optional_missing}


# -- WORD COUNT ----------------------------------------------------------------
def count_words(sections: dict) -> int:
    """
    Count total words across all memo body sections (excluding header fields).
    """
    body_fields = ["recommendation", "analysis", "risks", "conclusion", "additional_notes"]
    full_text = " ".join(sections.get(f, "") for f in body_fields)
    words = full_text.split()
    return len(words)


def build_warnings(word_count: int, validation: dict) -> list:
    """
    Build a list of warnings based on word count and missing fields.
    """
    warnings = []
    if word_count > WORD_COUNT_LIMIT:
        warnings.append(f"Word count ({word_count}) exceeds {WORD_COUNT_LIMIT} words. Consider condensing.")
    for field in ["recommendation", "conclusion", "subject"]:
        if field in validation["missing"]:
            warnings.append(f"'{field.title()}' field is missing from the memo.")
    return warnings


# -- WORD DOCUMENT CREATION ----------------------------------------------------
def add_bullet_or_paragraph(doc, text: str, style_name: str = "Normal"):
    """
    If a line starts with '- ' or '* ', add it as a bullet point.
    Otherwise add it as a normal paragraph.
    """
    lines = text.strip().splitlines()
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            p.runs[0].font.name = "Calibri"
            p.runs[0].font.size = Pt(11)
        else:
            p = doc.add_paragraph(line, style=style_name)
            if p.runs:
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.size = Pt(11)


def create_docx(sections: dict, output_path: str):
    """
    Create a professionally formatted Word document from parsed memo sections.
    Applies 1-inch margins, Calibri 11pt body, bold section headings,
    a memo header table, and bullet point conversion.
    """
    doc = Document()

    # -- Set 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # -- Title
    title_text = sections.get("title", "").strip() or "Business Memorandum"
    title_para = doc.add_paragraph(title_text)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(14)

    doc.add_paragraph()

    # -- Memo header table (To, From, Date, Subject)
    header_fields = [
        ("To", sections.get("to", "").strip()),
        ("From", sections.get("from", "").strip()),
        ("Date", sections.get("date", "").strip() or datetime.today().strftime("%B %d, %Y")),
        ("Subject", sections.get("subject", "").strip()),
    ]

    table = doc.add_table(rows=len(header_fields), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(header_fields):
        label_cell = table.cell(i, 0)
        value_cell = table.cell(i, 1)
        label_run = label_cell.paragraphs[0].add_run(label + ":")
        label_run.bold = True
        label_run.font.name = "Calibri"
        label_run.font.size = Pt(11)
        value_run = value_cell.paragraphs[0].add_run(value if value else "[Not provided]")
        value_run.font.name = "Calibri"
        value_run.font.size = Pt(11)

    doc.add_paragraph()

    # -- Body sections
    body_sections = [
        ("Recommendation", "recommendation"),
        ("Analysis", "analysis"),
        ("Risks", "risks"),
        ("Conclusion", "conclusion"),
        ("Additional Notes", "additional_notes"),
    ]

    for heading_text, field_key in body_sections:
        content = sections.get(field_key, "").strip()
        if not content:
            continue

        # Bold section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(heading_text)
        heading_run.bold = True
        heading_run.font.name = "Calibri"
        heading_run.font.size = Pt(12)

        # Section content with bullet detection
        add_bullet_or_paragraph(doc, content)
        doc.add_paragraph()

    doc.save(output_path)


# -- FORMATTING REPORT ---------------------------------------------------------
def create_report(
    input_path: str,
    output_docx: str,
    sections: dict,
    validation: dict,
    word_count: int,
    warnings: list,
    report_path: str
):
    """
    Write a Markdown formatting report summarizing what was applied,
    which fields were detected, warnings, and limitations.
    """
    found = validation["found"]
    missing = validation["missing"]
    optional_missing = validation["optional_missing"]

    lines = [
        "# Formatting Report",
        "",
        "## File Paths",
        f"- **Input file:** `{input_path}`",
        f"- **Output docx:** `{output_docx}`",
        f"- **Report:** `{report_path}`",
        "",
        "## Detected Fields",
    ]

    for field in REQUIRED_FIELDS:
        status = "Found" if field in found else "Missing"
        lines.append(f"- {field.title()}: {status}")

    for field in OPTIONAL_FIELDS:
        status = "Missing (optional)" if field in optional_missing else "Found"
        lines.append(f"- {field.title()}: {status}")

    lines += [
        "",
        "## Word Count",
        f"- Total body words: **{word_count}**",
    ]

    if warnings:
        lines += ["", "## Warnings"]
        for w in warnings:
            lines.append(f"- {w}")
    else:
        lines += ["", "## Warnings", "- None"]

    lines += [
        "",
        "## Formatting Applied",
        "- 1-inch margins",
        "- Calibri 11pt body font",
        "- Centered bold title",
        "- Memo header table with bold labels (To, From, Date, Subject)",
        "- Bold section headings (Recommendation, Analysis, Risks, Conclusion)",
        "- Bullet point conversion for lines starting with '- ' or '* '",
        "- 'Re:' and 'Regarding:' treated as equivalent to 'Subject:'",
        "- Unlabeled body text preserved under Additional Notes",
        "- Original content and meaning preserved",
        "",
        "## Limitations",
        "This script formats and checks structure only. It does not:",
        "- Judge whether the business argument is correct",
        "- Perform research or add external content",
        "- Invent or fill in missing fields",
        "- Guarantee a grade or evaluate against a rubric",
        "- Rewrite or edit the memo content",
    ]

    with open(report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


# -- MAIN ----------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Format a rough business memo draft into a professional Word document."
    )
    parser.add_argument("input_file", help="Path to the input .txt or .md memo file")
    parser.add_argument("--output-dir", default="outputs", help="Directory for output files (default: outputs)")
    args = parser.parse_args()

    input_path = args.input_file
    output_dir = args.output_dir

    # -- Validate input file
    if not os.path.exists(input_path):
        print(f"\nERROR: Input file not found: {input_path}\n")
        sys.exit(1)

    ext = os.path.splitext(input_path)[1].lower()
    if ext not in [".txt", ".md"]:
        print(f"\nERROR: Unsupported file type '{ext}'. Please provide a .txt or .md file.\n")
        sys.exit(1)

    with open(input_path, "r", encoding="utf-8") as f:
        raw_text = f.read()

    if not raw_text.strip():
        print(f"\nERROR: Input file is empty: {input_path}\n")
        sys.exit(1)

    # -- Create output directory
    os.makedirs(output_dir, exist_ok=True)

    output_docx = os.path.join(output_dir, "formatted_memo.docx")
    output_report = os.path.join(output_dir, "formatting_report.md")

    # -- Parse, validate, count
    print(f"Parsing memo from: {input_path}")
    sections = parse_memo(raw_text)

    print("Validating memo structure...")
    validation = validate_sections(sections)

    word_count = count_words(sections)
    warnings = build_warnings(word_count, validation)

    # -- Create Word document
    print(f"Creating Word document: {output_docx}")
    create_docx(sections, output_docx)

    # -- Create formatting report
    print(f"Writing formatting report: {output_report}")
    create_report(input_path, output_docx, sections, validation, word_count, warnings, output_report)

    # -- Summary
    print("\n" + "="*50)
    print("  Formatting complete.")
    print("="*50)
    print(f"  Output:  {output_docx}")
    print(f"  Report:  {output_report}")
    print(f"  Words:   {word_count}")
    if validation["missing"]:
        print(f"  Missing: {', '.join(validation['missing'])}")
    if warnings:
        print(f"  Warnings: {len(warnings)}")
    print("="*50 + "\n")


if __name__ == "__main__":
    main()
